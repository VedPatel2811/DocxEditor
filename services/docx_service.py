import io
import logging
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy

logger = logging.getLogger(__name__)


class SkillsSectionNotFoundError(Exception):
    pass


SKILLS_HEADINGS = {"skills", "technical skills", "core skills", "key skills"}


def add_skills_to_resume(file_bytes: bytes, skills: list[str]) -> bytes:
    doc = Document(io.BytesIO(file_bytes))

    # Find the Skills heading paragraph index
    skills_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() in SKILLS_HEADINGS:
            skills_idx = i
            logger.info(f"Skills section found at paragraph index {i}: '{para.text.strip()}'")
            break

    if skills_idx is None:
        searched = ", ".join(f'"{h}"' for h in SKILLS_HEADINGS)
        raise SkillsSectionNotFoundError(
            f"No Skills section heading found. Searched for: {searched}"
        )

    # Find the anchor: skip any separator paragraphs (empty lines, decorative
    # lines, Word horizontal rules) that sit between the heading and the skills.
    insert_after_idx = _find_insert_anchor(doc, skills_idx)
    logger.info(f"Inserting after paragraph index {insert_after_idx}")

    # Find an existing bullet paragraph in the section to use as a style template.
    # This ensures new bullets match the document's existing formatting exactly.
    template_para = _find_bullet_template(doc, skills_idx)
    logger.info(f"Bullet template style: '{template_para.style.name if template_para else None}'")

    # Insert in reverse so final order matches the input list
    for skill in reversed(skills):
        _insert_paragraph_after(doc, insert_after_idx, skill, template_para)
        logger.info(f"Inserted bullet: '{skill}'")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _find_insert_anchor(doc: Document, heading_idx: int) -> int:
    """Skip separator paragraphs after the heading and return the last one's index.

    A separator is: empty/whitespace-only, purely decorative chars (-_=),
    or contains a Word <w:hr> horizontal rule element.
    """
    anchor = heading_idx
    for i, para in enumerate(doc.paragraphs[heading_idx + 1:], start=heading_idx + 1):
        if para.style.name.lower().startswith("heading"):
            break
        text = para.text.strip()
        has_hr = para._element.find(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hr"
        ) is not None
        is_decorative = not text or all(c in "-_=• \t" for c in text)
        if has_hr or is_decorative:
            anchor = i
            continue
        break  # real content — stop
    return anchor


def _find_bullet_template(doc: Document, heading_idx: int):
    """Return the first real bullet/list paragraph after the heading, or None."""
    for para in doc.paragraphs[heading_idx + 1:]:
        if para.style.name.lower().startswith("heading"):
            break
        text = para.text.strip()
        style = para.style.name.lower()
        if text and ("bullet" in style or "list" in style):
            return para
    return None


def _insert_paragraph_after(doc: Document, after_idx: int, text: str, template_para):
    """Insert a new paragraph immediately after doc.paragraphs[after_idx].

    If a template_para (existing bullet) is available, deep-copy its XML so the
    new paragraph inherits identical formatting (style, indentation, numbering).
    Otherwise fall back to a fresh paragraph with List Bullet / plain bullet.
    """
    anchor_elem = doc.paragraphs[after_idx]._element

    if template_para is not None:
        # Copy the template bullet's full XML (preserves style + paragraph props)
        new_para_elem = deepcopy(template_para._element)
        # Remove all runs so we can set clean text
        for child in list(new_para_elem):
            if child.tag in (qn("w:r"), qn("w:hyperlink")):
                new_para_elem.remove(child)
    else:
        # No existing bullet found — create a plain paragraph element
        new_para_elem = OxmlElement("w:p")

    anchor_elem.addnext(new_para_elem)

    # Locate the python-docx wrapper for the newly inserted element
    new_paragraph = next((p for p in doc.paragraphs if p._element is new_para_elem), None)
    if new_paragraph is None:
        return

    if template_para is not None:
        new_paragraph.text = text
    else:
        _apply_list_bullet_or_plain(doc, new_paragraph, text)


def _apply_list_bullet_or_plain(doc: Document, paragraph, text: str):
    """Try 'List Bullet' style; fall back to plain '• ' prefix."""
    try:
        paragraph.style = doc.styles["List Bullet"]
        paragraph.text = text
    except KeyError:
        paragraph.style = doc.styles["Normal"]
        paragraph.text = f"• {text}"
